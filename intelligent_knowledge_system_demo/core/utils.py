"""
================================================================================
智识通 - 企业智能知识库系统
工具函数库 (function_tools.py)

【这个文件的作用】
这个文件是系统的"工具箱"，提供了各种基础功能：
1. 向量数据库操作（存数据、搜数据）
2. 文档读取（读取Word文档）
3. AI模型调用（调用通义千问等大模型）
4. 工具函数（中文转拼音等）

【什么是向量（Vector）】
向量是一串数字，可以表示文本的含义。
比如："苹果" 和 "水果" 的向量会很相似，因为它们意思相关。
这样计算机就能"理解"文本的相似度了。

【什么是Embedding】
Embedding就是把文字转换成向量的过程。
我们使用阿里云的text-embedding模型来完成这个转换。
================================================================================
"""

# 导入需要的库
from functools import wraps  # 装饰器工具
from pypinyin import pinyin, Style  # 中文转拼音
from docx import Document  # 读取Word文档

from langchain_text_splitters import RecursiveCharacterTextSplitter  # 文本分割工具

from models import *  # 从models.py导入所有模型配置
from  MyVectorDBConnector import *

# =============================================================================
# 第二部分：文档处理函数
# =============================================================================

def extract_text_from_docx(filename):
    """
    【功能】从Word文档中提取文字内容
    
    【参数】
        filename: Word文档的文件路径
    
    【返回值】
        文档内容列表，每个元素是一段文字（已经被切分）
    
    【处理流程】
        1. 打开Word文档
        2. 读取所有段落的文字
        3. 将长文本切分成小块（方便检索）
    
    【为什么要切分】
    一篇文档可能很长（几千字），如果整体存储：
    - 检索时只能整篇匹配，不够精确
    - 可能超出模型处理长度限制
    
    切分成小块后：
    - 可以精确匹配到具体段落
    - 检索结果更精准
    """
    print(f'【文档读取】正在读取: {filename}')
    
    full_text = ''  # 存储完整文本
    
    # Document() 打开Word文档
    doc = Document(filename)
    
    # 遍历文档中的所有段落
    # doc.paragraphs 是段落列表
    for para in doc.paragraphs:
        # para.text 获取段落文字
        # .strip() 去除首尾空白
        # 只保留非空段落
        if para.text.strip():
            full_text += para.text + '\n'  # 加换行符分隔段落
    
    print(f'【文档读取】原文共 {len(full_text)} 个字符')
    
    # 使用RecursiveCharacterTextSplitter切分文本
    # chunk_size=300: 每个块大约300个字符
    # chunk_overlap=30: 相邻块重叠30个字符（避免信息断裂）
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=300,  # 块大小
        chunk_overlap=30  # 重叠大小
    )
    
    # split_text() 执行切分
    documents = splitter.split_text(full_text)
    
    print(f'【文档读取】切分成 {len(documents)} 个片段')
    
    return documents


# =============================================================================
# 第三部分：AI模型调用函数
# =============================================================================

def get_completion(prompt, model=ALI_TONGYI_TURBO_MODEL):
    """
    【功能】调用大语言模型生成回答
    
    【参数】
        prompt: 提示词（给AI的指令和问题）
        model: 使用的模型名称，默认是阿里云的qwen3.5-flash
    
    【返回值】
        AI生成的文本回答
    
    【什么是Prompt（提示词）】
    告诉AI要做什么的指令。比如：
    "你是一个问答助手。根据以下资料回答问题：..."
    
    【temperature参数】
    控制AI输出的随机性：
    - 0: 最确定，每次回答都一样（适合问答）
    - 1: 最随机，每次回答都不同（适合创作）
    """
    # 构建消息格式（OpenAI标准格式）
    # role="user" 表示这是用户的消息
    messages = [{"role": "user", "content": prompt}]
    
    # 获取AI客户端
    client = get_normal_client()
    
    # 调用API生成回答
    # chat.completions.create() 是聊天API
    # model: 使用的模型
    # messages: 对话历史（这里只有一条）
    # temperature=0: 输出最确定性的结果
    response = client.chat.completions.create(
        model=model,
        messages=messages,
        temperature=0,  # 0表示最确定，不随机
    )
    
    # 从响应中提取AI的回答文本
    # response.choices[0] 是第一个回答选项
    # .message.content 是回答的内容
    return response.choices[0].message.content


# =============================================================================
# 第四部分：工具装饰器
# =============================================================================

def to_pinyin(fn):
    """
    【功能】装饰器：将中文集合名转换为拼音
    
    【作用】
    有些数据库不支持中文作为标识符，所以把中文转成拼音更安全。
    
    【例子】
    "人事管理流程.docx" -> "renshiguanliliuchengdocx"
    
    【使用方式】
    @to_pinyin
    def my_function(collection_name='demo'):
        ...
    
    调用 my_function(collection_name="测试") 时，
    实际传入的是 "ceshi"
    """
    @wraps(fn)  # 保留原函数的元信息
    def chinese_to_pinyin(*args, **kwargs):
        # 获取collection_name参数
        chinese_name = kwargs['collection_name']
        
        # 去掉文件名中的点号（.）
        chinese_name = chinese_name.replace('.', '')
        
        # 使用pypinyin库将中文转为拼音
        # style=Style.NORMAL: 普通风格，不带声调
        # heteronym=False: 不使用多音字
        pinyin_list = pinyin(chinese_name, style=Style.NORMAL, heteronym=False)
        
        # 将拼音列表拼接成字符串
        # word[0] 取每个字的第一个拼音
        # .lower() 转成小写
        pinyin_str = ''.join([word[0].lower() for word in pinyin_list])
        
        # 替换参数中的中文名为拼音
        kwargs['collection_name'] = pinyin_str
        
        # 调用原函数
        return fn(*args, **kwargs)
    
    return chinese_to_pinyin
